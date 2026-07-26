"""Testes de escrever_saidas (nucleo.escrita) usando um resultado sintético."""

import json
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from nucleo.escrita import escrever_saidas
from nucleo.formatacao import ResultadoTranscricao, Segmento


def _resultado(tmp_path):
    return ResultadoTranscricao(
        arquivo=tmp_path / "meu_video.mp4",
        idioma="pt",
        probabilidade_idioma=0.99,
        duracao=5.0,
        segmentos=[
            Segmento(inicio=0.0, fim=2.5, texto="Primeira linha."),
            Segmento(inicio=2.5, fim=5.0, texto="Segunda linha."),
        ],
    )


def test_gera_txt(tmp_path):
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["txt"])
    assert len(gerados) == 1
    conteudo = gerados[0].read_text(encoding="utf-8")
    assert conteudo == "Primeira linha.\nSegunda linha.\n"


def test_gera_srt(tmp_path):
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["srt"])
    conteudo = gerados[0].read_text(encoding="utf-8")
    assert "1\n00:00:00,000 --> 00:00:02,500\nPrimeira linha." in conteudo
    assert "2\n00:00:02,500 --> 00:00:05,000\nSegunda linha." in conteudo


def test_gera_vtt(tmp_path):
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["vtt"])
    conteudo = gerados[0].read_text(encoding="utf-8")
    assert conteudo.startswith("WEBVTT")
    assert "00:00:00.000 --> 00:00:02.500" in conteudo


def test_gera_json(tmp_path):
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["json"])
    dados = json.loads(gerados[0].read_text(encoding="utf-8"))
    assert dados["idioma"] == "pt"
    assert len(dados["segmentos"]) == 2
    assert dados["segmentos"][0]["texto"] == "Primeira linha."


def test_gera_todos_os_formatos(tmp_path):
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["txt", "srt", "vtt", "json"])
    extensoes = sorted(g.suffix for g in gerados)
    assert extensoes == [".json", ".srt", ".txt", ".vtt"]


def test_gera_html(tmp_path):
    """HTML sai sempre que pedido; independe de libs externas."""
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["html"])
    assert len(gerados) == 1
    assert gerados[0].suffix == ".html"
    corpo = gerados[0].read_text(encoding="utf-8")
    assert "<!DOCTYPE html>" in corpo
    assert "Primeira linha." in corpo
    assert "Segunda linha." in corpo
    assert 'lang="pt-BR"' in corpo


def _resultado_com_falantes(tmp_path):
    from nucleo.formatacao import ResultadoTranscricao, Segmento
    return ResultadoTranscricao(
        arquivo=tmp_path / "entrevista.mp4",
        idioma="pt",
        probabilidade_idioma=0.99,
        duracao=6.0,
        segmentos=[
            Segmento(0.0, 2.0, "Bom dia.", falante="FALANTE_01"),
            Segmento(2.0, 4.0, "Como vai?", falante="FALANTE_01"),
            Segmento(4.0, 6.0, "Tudo bem.", falante="FALANTE_02"),
        ],
        falantes=["FALANTE_01", "FALANTE_02"],
    )


def test_html_com_falantes(tmp_path):
    """Quando há diarização, o HTML marca cada fala com nome + carimbo de tempo."""
    res = _resultado_com_falantes(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["html"])
    corpo = gerados[0].read_text(encoding="utf-8")
    assert "Falante 1" in corpo
    assert "Falante 2" in corpo
    assert "<h2>Falantes</h2>" in corpo


def test_gera_html_junto_com_txt(tmp_path):
    """Pedir HTML junto com TXT não muda o TXT e ainda gera o HTML."""
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["txt", "html"])
    extensoes = sorted(g.suffix for g in gerados)
    assert ".html" in extensoes
    assert ".txt" in extensoes
    txt_gerado = [g for g in gerados if g.suffix == ".txt"][0]
    assert txt_gerado.read_text(encoding="utf-8") == "Primeira linha.\nSegunda linha.\n"


def test_gera_docx(tmp_path):
    """DOCX sai quando a dep python-docx está instalada."""
    pytest.importorskip("docx")
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["docx"])
    assert len(gerados) == 1
    assert gerados[0].suffix == ".docx"
    # Assinatura DOCX (ZIP): os primeiros 4 bytes precisam ser PK\x03\x04
    cabecalho = gerados[0].read_bytes()[:4]
    assert cabecalho.startswith(b"PK")
    # Reabre com python-docx e confere que o texto da transcrição está lá.
    from docx import Document  # type: ignore
    doc = Document(gerados[0])
    paragrafos = "\n".join(p.text for p in doc.paragraphs)
    assert "Primeira linha." in paragrafos
    assert "Segunda linha." in paragrafos


def test_docx_com_falantes(tmp_path):
    """DOCX rotula cada bloco com nome do falante quando há diarização."""
    pytest.importorskip("docx")
    res = _resultado_com_falantes(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["docx"])
    from docx import Document  # type: ignore
    doc = Document(gerados[0])
    paragrafos = "\n".join(p.text for p in doc.paragraphs)
    assert "Falante 1" in paragrafos
    assert "Falante 2" in paragrafos


def test_gera_pdf(tmp_path):
    """PDF sai quando a dep fpdf2 está instalada."""
    pytest.importorskip("fpdf")
    res = _resultado(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["pdf"])
    assert len(gerados) == 1
    assert gerados[0].suffix == ".pdf"
    # Assinatura PDF: os primeiros 5 bytes precisam ser %PDF-
    cabecalho = gerados[0].read_bytes()[:5]
    assert cabecalho.startswith(b"%PDF-")


def test_pdf_com_falantes(tmp_path):
    """PDF inclui o nome dos falantes quando a transcrição é diarizada."""
    pytest.importorskip("fpdf")
    res = _resultado_com_falantes(tmp_path)
    gerados = escrever_saidas(res, tmp_path / "out", ["pdf"])
    # PDF é binário — conferimos só que gerou e não está vazio.
    assert gerados[0].stat().st_size > 500


def test_docx_sem_dep_levanta_erro(tmp_path, monkeypatch):
    """Sem python-docx, pedir .docx tem que falhar com mensagem clara."""
    import builtins
    original = builtins.__import__
    def bloquear_docx(name, *args, **kwargs):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("docx bloqueado pelo teste")
        return original(name, *args, **kwargs)
    monkeypatch.setattr(builtins, "__import__", bloquear_docx)
    res = _resultado(tmp_path)
    # Remove do cache para forçar reimport
    import sys
    sys.modules.pop("docx", None)
    with pytest.raises(RuntimeError, match="python-docx"):
        escrever_saidas(res, tmp_path / "out", ["docx"])


def test_pdf_sem_dep_levanta_erro(tmp_path, monkeypatch):
    """Sem fpdf2, pedir .pdf tem que falhar com mensagem clara."""
    import builtins
    original = builtins.__import__
    def bloquear_fpdf(name, *args, **kwargs):
        if name == "fpdf" or name.startswith("fpdf."):
            raise ImportError("fpdf bloqueado pelo teste")
        return original(name, *args, **kwargs)
    monkeypatch.setattr(builtins, "__import__", bloquear_fpdf)
    res = _resultado(tmp_path)
    import sys
    sys.modules.pop("fpdf", None)
    with pytest.raises(RuntimeError, match="fpdf2"):
        escrever_saidas(res, tmp_path / "out", ["pdf"])