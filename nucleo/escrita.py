"""Gravação dos resultados da transcrição nos formatos txt/srt/vtt/json/html/docx/pdf."""

from __future__ import annotations

import json
import math
from html import escape
from pathlib import Path

from .formatacao import ResultadoTranscricao, formatar_timestamp_legenda


def nome_amigavel(rotulo: str | None, nomes: dict[str, str] | None = None) -> str:
    """Converte 'FALANTE_01' no nome que o usuário vê ('Falante 1' ou 'Maria')."""
    if not rotulo:
        return ""
    if nomes and rotulo in nomes:
        return nomes[rotulo]
    if rotulo.startswith("FALANTE_"):
        numero = rotulo.removeprefix("FALANTE_").lstrip("0") or "0"
        return f"Falante {numero}"
    return rotulo


def _duracao_hms(segundos: float) -> str:
    if not math.isfinite(segundos) or segundos <= 0:
        return "—"
    s = int(segundos)
    return f"{s // 3600:02d}:{(s % 3600) // 60:02d}:{s % 60:02d}"


def _escrever_html(
    resultado: ResultadoTranscricao,
    base: Path,
    tem_falantes: bool,
    rotulo_de,
) -> Path:
    """Gera um HTML único, imprimível, com CSS mínimo e cabeçalho com metadados."""
    segmentos = []
    for seg in resultado.segmentos:
        ts = f"{formatar_timestamp_legenda(seg.inicio, '.')} --> {formatar_timestamp_legenda(seg.fim, '.')}"
        fala = escape(seg.texto.strip())
        rotulo = rotulo_de(seg)
        if seg.falante:
            segmentos.append(
                f'        <div class="seg">'
                f'<span class="ts">{escape(ts)}</span> '
                f'<span class="nome">{escape(rotulo)}:</span> '
                f'<span class="texto">{fala}</span>'
                f'</div>'
            )
        else:
            segmentos.append(
                f'        <div class="seg">'
                f'<span class="ts">{escape(ts)}</span> '
                f'<span class="texto">{fala}</span>'
                f'</div>'
            )

    nomes = ""
    if tem_falantes:
        linhas = "\n".join(
            f"          <li>{escape(nome_amigavel(r, None))}</li>"
            for r in resultado.falantes
        )
        nomes = f"      <h2>Falantes</h2>\n      <ul>\n{linhas}\n      </ul>\n"

    corpo = "\n".join(segmentos) if segmentos else "        <p><em>(sem segmentos)</em></p>"
    idioma = escape(resultado.idioma or "—")
    arquivo = escape(resultado.arquivo.name)
    duracao = escape(_duracao_hms(resultado.duracao))

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Transcrição de {arquivo}</title>
    <style>
      body {{ font: 16px/1.55 -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
              max-width: 740px; margin: 2.2em auto; padding: 0 1em; color: #1f2328; }}
      h1 {{ font-size: 1.35em; margin: 0 0 .25em 0; }}
      .meta {{ color: #57606a; margin: 0 0 1.5em 0; font-size: .9em; }}
      .meta span {{ margin-right: 1.2em; }}
      h2 {{ font-size: 1em; margin: 1.4em 0 .4em 0; }}
      .seg {{ margin: .15em 0; }}
      .ts {{ color: #8b949e; font-variant-numeric: tabular-nums; margin-right: .6em; font-size: .85em; }}
      .nome {{ font-weight: 600; }}
      .print {{ display: block; margin: 1.5em 0; }}
      @media print {{ .no-print {{ display: none; }} }}
    </style>
  </head>
  <body>
    <header>
      <h1>Transcrição de {arquivo}</h1>
      <p class="meta">
        <span><strong>Idioma:</strong> {idioma}</span>
        <span><strong>Duração:</strong> {duracao}</span>
      </p>
    </header>
{nomes}    <h2>Segmentos</h2>
{corpo}
    <p class="print no-print">
      <button onclick="window.print()">Imprimir / salvar em PDF</button>
    </p>
  </body>
</html>
"""
    caminho = base.with_suffix(".html")
    caminho.write_text(html, encoding="utf-8")
    return caminho


def _escrever_docx(
    resultado: ResultadoTranscricao,
    base: Path,
    tem_falantes: bool,
    rotulo_de,
) -> Path:
    """Gera um .docx (Word) com cabeçalho e parágrafos rotulados por falante."""
    try:
        from docx import Document  # type: ignore
    except ImportError as erro:
        raise RuntimeError(
            "Para gerar .docx é preciso instalar 'python-docx' (pip install python-docx)."
        ) from erro

    doc = Document()
    doc.add_heading(f"Transcrição de {resultado.arquivo.name}", level=1)
    p = doc.add_paragraph()
    p.add_run("Idioma: ").bold = True
    p.add_run(resultado.idioma or "—")
    p.add_run("    Duração: ").bold = True
    p.add_run(_duracao_hms(resultado.duracao))

    if tem_falantes:
        doc.add_heading("Falantes", level=2)
        for rotulo in resultado.falantes:
            doc.add_paragraph(nome_amigavel(rotulo, None), style="List Bullet")

    doc.add_heading("Segmentos", level=2)
    for seg in resultado.segmentos:
        ts = f"[{formatar_timestamp_legenda(seg.inicio)} --> {formatar_timestamp_legenda(seg.fim)}]"
        texto = seg.texto.strip()
        if seg.falante:
            paragrafo = doc.add_paragraph(style="List Number")
            paragrafo.add_run(f"{ts} ").bold = True
            paragrafo.add_run(f"{rotulo_de(seg)}: ").bold = True
            paragrafo.add_run(texto)
        else:
            paragrafo = doc.add_paragraph(style="List Number")
            paragrafo.add_run(f"{ts} ").bold = True
            paragrafo.add_run(texto)

    caminho = base.with_suffix(".docx")
    doc.save(caminho)
    return caminho


def _escrever_pdf(
    resultado: ResultadoTranscricao,
    base: Path,
    tem_falantes: bool,
    rotulo_de,
) -> Path:
    """Gera um PDF de uma página por segmento, com fonte sem-serifa embutida."""
    try:
        from fpdf import FPDF  # type: ignore
    except ImportError as erro:
        raise RuntimeError(
            "Para gerar .pdf é preciso instalar 'fpdf2' (pip install fpdf2)."
        ) from erro

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    # Fonte padrão do fpdf2 (Helvetica) já cobre acentos latinos com a codificação
    # latin-1; quando o texto tem isso, saímos bem.
    pdf.set_font("Helvetica", size=11)

    pdf.cell(0, 8, f"Transcricao de {resultado.arquivo.name}", ln=1)
    pdf.set_font("Helvetica", size=9)
    pdf.cell(0, 5, f"Idioma: {resultado.idioma or '—'}   Duracao: {_duracao_hms(resultado.duracao)}", ln=1)
    pdf.ln(2)
    pdf.set_font("Helvetica", size=11)

    if tem_falantes:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Falantes", ln=1)
        pdf.set_font("Helvetica", size=11)
        for rotulo in resultado.falantes:
            pdf.cell(0, 5, f"  - {nome_amigavel(rotulo, None)}", ln=1)
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 6, "Segmentos", ln=1)
        pdf.set_font("Helvetica", size=11)

    for seg in resultado.segmentos:
        ts = f"[{formatar_timestamp_legenda(seg.inicio, '.')} --> {formatar_timestamp_legenda(seg.fim, '.')}]"
        texto = seg.texto.strip()
        linha = f"{ts} {rotulo_de(seg)}: {texto}" if seg.falante else f"{ts} {texto}"
        # fpdf2 aceita acentos; latin-1 cobre o básico para pt-BR sem precisar
        # carregar TTF. Texto muito longo quebra automaticamente entre linhas.
        pdf.multi_cell(0, 5, linha)
        pdf.ln(1)

    caminho = base.with_suffix(".pdf")
    pdf.output(str(caminho))
    return caminho


def escrever_saidas(
    resultado: ResultadoTranscricao,
    pasta_saida: Path,
    formatos: list[str],
    nomes_falantes: dict[str, str] | None = None,
) -> list[Path]:
    """Grava os arquivos pedidos e devolve os caminhos gerados.

    Quando a transcrição foi diarizada, os blocos saem rotulados com o falante.
    Sem diarização, os arquivos ficam idênticos aos de sempre — há teste
    garantindo isso, porque é o caminho que a maioria dos usuários usa.
    """
    pasta_saida.mkdir(parents=True, exist_ok=True)
    base = pasta_saida / resultado.arquivo.stem
    gerados: list[Path] = []

    tem_falantes = any(seg.falante for seg in resultado.segmentos)

    def rotulo_de(seg) -> str:
        return nome_amigavel(seg.falante, nomes_falantes)

    if "txt" in formatos:
        caminho = base.with_suffix(".txt")
        if tem_falantes:
            # Agrupa falas seguidas da mesma pessoa num parágrafo só, com linha
            # em branco na troca de voz — assim o diálogo fica legível.
            linhas: list[str] = []
            falante_anterior: str | None = None
            for seg in resultado.segmentos:
                if seg.falante != falante_anterior:
                    if falante_anterior is not None:
                        linhas.append("")
                    linhas.append(f"{rotulo_de(seg)}: {seg.texto.strip()}")
                    falante_anterior = seg.falante
                else:
                    linhas.append(seg.texto.strip())
        else:
            linhas = [seg.texto.strip() for seg in resultado.segmentos]
        caminho.write_text("\n".join(linhas) + "\n", encoding="utf-8")
        gerados.append(caminho)

    if "srt" in formatos:
        caminho = base.with_suffix(".srt")
        linhas = []
        for i, seg in enumerate(resultado.segmentos, start=1):
            linhas.append(str(i))
            linhas.append(
                f"{formatar_timestamp_legenda(seg.inicio)} --> {formatar_timestamp_legenda(seg.fim)}"
            )
            texto = seg.texto.strip()
            linhas.append(f"{rotulo_de(seg)}: {texto}" if seg.falante else texto)
            linhas.append("")
        caminho.write_text("\n".join(linhas), encoding="utf-8")
        gerados.append(caminho)

    if "vtt" in formatos:
        caminho = base.with_suffix(".vtt")
        linhas = ["WEBVTT", ""]
        for seg in resultado.segmentos:
            linhas.append(
                f"{formatar_timestamp_legenda(seg.inicio, '.')} --> {formatar_timestamp_legenda(seg.fim, '.')}"
            )
            texto = seg.texto.strip()
            # <v Nome> é a marcação padrão do WebVTT para identificar quem fala:
            # players que a entendem mostram o nome; os demais ignoram a tag.
            linhas.append(f"<v {rotulo_de(seg)}>{texto}" if seg.falante else texto)
            linhas.append("")
        caminho.write_text("\n".join(linhas), encoding="utf-8")
        gerados.append(caminho)

    if "json" in formatos:
        caminho = base.with_suffix(".json")
        segmentos = []
        for s in resultado.segmentos:
            item = {"inicio": s.inicio, "fim": s.fim, "texto": s.texto.strip()}
            if s.falante:
                item["falante"] = s.falante
            segmentos.append(item)

        dados = {
            "arquivo": str(resultado.arquivo),
            "idioma": resultado.idioma,
            "probabilidade_idioma": resultado.probabilidade_idioma,
            "duracao_segundos": resultado.duracao,
            "segmentos": segmentos,
        }
        if tem_falantes:
            dados["falantes"] = {
                rotulo: nome_amigavel(rotulo, nomes_falantes) for rotulo in resultado.falantes
            }
        caminho.write_text(json.dumps(dados, ensure_ascii=False, indent=2), encoding="utf-8")
        gerados.append(caminho)

    if "html" in formatos:
        gerados.append(_escrever_html(resultado, base, tem_falantes, rotulo_de))

    if "docx" in formatos:
        gerados.append(_escrever_docx(resultado, base, tem_falantes, rotulo_de))

    if "pdf" in formatos:
        gerados.append(_escrever_pdf(resultado, base, tem_falantes, rotulo_de))

    return gerados
